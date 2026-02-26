import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def apply_formal_formatting(doc):
    """Applies Arial, 11pt, 1.5 spacing, and Justified alignment."""
    
    # 1. Update the baseline 'Normal' style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # 2. Iterate through paragraphs to apply specific rules
    for para in doc.paragraphs:
        # Standard Text
        if para.style.name == 'Normal':
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_after = Pt(12)
        
        # Headings (Bold, Arial, Left-Aligned)
        elif para.style.name.startswith('Heading'):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.bold = True

def format_existing_file():
    print("\n--- 📂 Format Existing Document ---")
    file_path = input("Enter the path of the DOCX file: ").strip('"')
    
    if not os.path.exists(file_path) or not file_path.endswith('.docx'):
        print("❌ Error: Valid DOCX file not found!")
        return

    try:
        doc = Document(file_path)
        print("🔄 Applying formatting...")
        apply_formal_formatting(doc)
        
        output_path = file_path.replace(".docx", "_formatted.docx")
        doc.save(output_path)
        print(f"✅ Success! Saved as: {output_path}")

    except Exception as e:
        print(f"❌ Error: {e}")

def create_from_raw_text():
    print("\n--- 📝 Create from Raw Text ---")
    print("Paste your text below. Type 'DONE' on a new line when finished:")
    
    lines = []
    while True:
        line = input()
        if line.strip().upper() == 'DONE':
            break
        lines.append(line)
    
    if not lines:
        print("❌ No text entered!")
        return

    filename = input("\nEnter name for new file (e.g., report): ")
    if not filename.endswith('.docx'):
        filename += ".docx"
        
    try:
        doc = Document()
        # Add the text as a single paragraph (or split by double newlines if you prefer)
        text_block = "\n".join(lines)
        doc.add_paragraph(text_block)
        
        print("🔄 Formatting and saving...")
        apply_formal_formatting(doc)
        
        doc.save(filename)
        print(f"✅ Success! Created and saved as: {filename}")

    except Exception as e:
        print(f"❌ Error: {e}")

def main_menu():
    while True:
        print("\n" + "="*40)
        print("   SCRIPT SAVANT: DOCUMENT FORMATTER")
        print("="*40)
        print("1. 📂 Clean up an Existing DOCX File")
        print("2. 📝 Create New DOCX from Raw Text")
        print("0. 🚪 Exit")
        
        choice = input("\nEnter choice (0-2): ")
        
        if choice == '1': format_existing_file()
        elif choice == '2': create_from_raw_text()
        elif choice == '0': break
        else: print("Invalid choice!")

if __name__ == "__main__":
    main_menu()
